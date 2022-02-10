def getFiles(path):
    '''
    It takes path as parameter
    Returns list of files present in the path
    '''

    try:
        import os
        allFile=[]
        for filepath, folder, files in os.walk(path):
            for file in files:
                allFile.append(filepath+"\\"+file)
        return allFile

    except Exception as e:
        return e


def mergePDF(path, mergedFilename):
    '''
    It takes path, mergedFilename as parameter
    It merges pdf files
    '''
    try:
        from PyPDF2 import PdfFileMerger
        mergerObj=PdfFileMerger()
        pdfFiles=[]
        for file in getFiles(path):
            if file.endswith('.pdf'):
                pdfFiles.append(file)
        if len(pdfFiles)==0:
            return 'No Pdf Files found!'
        elif len(pdfFiles)==1:
            return 'only single pdf file is present!'
        else:
            for file in pdfFiles:
                mergerObj.append(file)
            mergerObj.write(f"{mergedFilename}.pdf")
            return f'{len(pdfFiles)} Pdf files merged!'

    except Exception as e:
        return e


def mergeDOC(path, mergedFilename):
    '''
    It takes path,mergedFilename as parameter
    It merges .docx file
    '''
    try:
        from docx import Document
        mergedFile=Document()
        docFiles=[]
        for file in getFiles(path):
            if file.endswith('.docx'):
                docFiles.append(file)
        if len(docFiles)==0:
            return 'No docx Files found!'
        elif len(docFiles)==1:
            return 'only single docx file is present!'
        else:
            for index, file in enumerate(docFiles):
                doc=Document(file)

                for paragraph in doc.paragraphs:
                    mergedFile.add_paragraph(paragraph.text)
                if index < len(docFiles) - 1:
                    mergedFile.add_page_break()
            mergedFile.save(f'{mergedFilename}.docx')
            return f"{len(docFiles)} docx files merged!"

    except Exception as e:
        return e


def mergeTXT(path, mergedFilename):
    '''
    It takes path,mergedFilename as parameter
    It merges .txt file
    '''
    try:
        txtFiles=[]
        for file in getFiles(path):
            if file.endswith('.txt'):
                txtFiles.append(file)
        if len(txtFiles)==0:
            return 'No txt Files found!'
        elif len(txtFiles)==1:
            return 'only single txt file is present!'
        else:
            with open(f'{mergedFilename}.txt', 'w') as finalFile:
                for file in txtFiles:
                    with open(file, 'r') as f:
                        finalFile.write(f.read()+'\n')
            return f'{len(txtFiles)} files merged!'

    except Exception as e:
        return e

